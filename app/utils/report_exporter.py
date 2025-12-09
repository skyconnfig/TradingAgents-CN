"""
æŠ¥å‘Šå¯¼å‡ºå·¥å…· - æ”¯æŒ Markdownã€Wordã€PDF æ ¼å¼

ä¾èµ–å®‰è£…:
    pip install pypandoc markdown

PDF å¯¼å‡ºéœ€è¦é¢å¤–å·¥å…·:
    - wkhtmltopdf (æ¨è): https://wkhtmltopdf.org/downloads.html
    - æˆ– LaTeX: https://www.latex-project.org/get/
"""

import logging
import os
import tempfile
from pathlib import Path
from typing import Dict, Any, Optional

logger = logging.getLogger(__name__)

# æ£€æŸ¥ä¾èµ–æ˜¯å¦å¯ç”¨
try:
    import markdown
    import pypandoc

    # æ£€æŸ¥ pandoc æ˜¯å¦å¯ç”¨
    try:
        pypandoc.get_pandoc_version()
        PANDOC_AVAILABLE = True
        logger.info("âœ… Pandoc å¯ç”¨")
    except OSError:
        PANDOC_AVAILABLE = False
        logger.warning("âš ï¸ Pandoc ä¸å¯ç”¨ï¼ŒWord å’Œ PDF å¯¼å‡ºåŠŸèƒ½å°†ä¸å¯ç”¨")

    EXPORT_AVAILABLE = True
except ImportError as e:
    EXPORT_AVAILABLE = False
    PANDOC_AVAILABLE = False
    logger.warning(f"âš ï¸ å¯¼å‡ºåŠŸèƒ½ä¾èµ–åŒ…ç¼ºå¤±: {e}")
    logger.info("ğŸ’¡ è¯·å®‰è£…: pip install pypandoc markdown")

# æ£€æŸ¥ pdfkitï¼ˆå”¯ä¸€çš„ PDF ç”Ÿæˆå·¥å…·ï¼‰
PDFKIT_AVAILABLE = False
PDFKIT_ERROR = None

try:
    import pdfkit
    # æ£€æŸ¥ wkhtmltopdf æ˜¯å¦å®‰è£…
    try:
        pdfkit.configuration()
        PDFKIT_AVAILABLE = True
        logger.info("âœ… pdfkit + wkhtmltopdf å¯ç”¨ï¼ˆPDF ç”Ÿæˆå·¥å…·ï¼‰")
    except Exception as e:
        PDFKIT_ERROR = str(e)
        logger.warning("âš ï¸ wkhtmltopdf æœªå®‰è£…ï¼ŒPDF å¯¼å‡ºåŠŸèƒ½ä¸å¯ç”¨")
        logger.info("ğŸ’¡ å®‰è£…æ–¹æ³•: https://wkhtmltopdf.org/downloads.html")
except ImportError:
    logger.warning("âš ï¸ pdfkit æœªå®‰è£…ï¼ŒPDF å¯¼å‡ºåŠŸèƒ½ä¸å¯ç”¨")
    logger.info("ğŸ’¡ å®‰è£…æ–¹æ³•: pip install pdfkit")
except Exception as e:
    PDFKIT_ERROR = str(e)
    logger.warning(f"âš ï¸ pdfkit æ£€æµ‹å¤±è´¥: {e}")

# æ£€æŸ¥ xhtml2pdf (çº¯ Python PDF ç”Ÿæˆå·¥å…·)
XHTML2PDF_AVAILABLE = False
try:
    from xhtml2pdf import pisa
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    
    # å°è¯•æ³¨å†Œä¸­æ–‡å­—ä½“ (Windows å¸¸ç”¨ SimHei)
    SIMHEI_PATH = r'C:\Windows\Fonts\simhei.ttf'
    if os.path.exists(SIMHEI_PATH):
        pdfmetrics.registerFont(TTFont('SimHei', SIMHEI_PATH))
        HAS_CHINESE_FONT = True
        logger.info("âœ… xhtml2pdf å¯ç”¨ (å·²åŠ è½½ SimHei å­—ä½“)")
    else:
        HAS_CHINESE_FONT = False
        logger.warning("âš ï¸ xhtml2pdf: æœªæ‰¾åˆ° SimHei å­—ä½“ï¼Œä¸­æ–‡å¯èƒ½ä¹±ç ")
        
    XHTML2PDF_AVAILABLE = True
except ImportError:
    logger.warning("âš ï¸ xhtml2pdf æœªå®‰è£…")
except Exception as e:
    logger.warning(f"âš ï¸ xhtml2pdf åˆå§‹åŒ–å¤±è´¥: {e}")


class ReportExporter:
    """æŠ¥å‘Šå¯¼å‡ºå™¨ - æ”¯æŒ Markdownã€Wordã€PDF æ ¼å¼"""

    def __init__(self):
        self.export_available = EXPORT_AVAILABLE
        self.pandoc_available = PANDOC_AVAILABLE
        self.export_available = EXPORT_AVAILABLE
        self.pandoc_available = PANDOC_AVAILABLE
        self.pdfkit_available = PDFKIT_AVAILABLE
        self.xhtml2pdf_available = XHTML2PDF_AVAILABLE

        logger.info("ğŸ“‹ ReportExporter åˆå§‹åŒ–:")
        logger.info(f"  - export_available: {self.export_available}")
        logger.info(f"  - pandoc_available: {self.pandoc_available}")
        logger.info(f"  - pandoc_available: {self.pandoc_available}")
        logger.info(f"  - pdfkit_available: {self.pdfkit_available}")
        logger.info(f"  - xhtml2pdf_available: {self.xhtml2pdf_available}")
    
    def generate_markdown_report(self, report_doc: Dict[str, Any]) -> str:
        """ç”Ÿæˆ Markdown æ ¼å¼æŠ¥å‘Š"""
        logger.info("ğŸ“ ç”Ÿæˆ Markdown æŠ¥å‘Š...")
        
        stock_symbol = report_doc.get("stock_symbol", "unknown")
        analysis_date = report_doc.get("analysis_date", "")
        analysts = report_doc.get("analysts", [])
        research_depth = report_doc.get("research_depth", 1)
        reports = report_doc.get("reports", {})
        summary = report_doc.get("summary", "")
        
        content_parts = []
        
        # æ ‡é¢˜å’Œå…ƒä¿¡æ¯
        content_parts.append(f"# {stock_symbol} è‚¡ç¥¨åˆ†ææŠ¥å‘Š")
        content_parts.append("")
        content_parts.append(f"**åˆ†ææ—¥æœŸ**: {analysis_date}")
        if analysts:
            content_parts.append(f"**åˆ†æå¸ˆ**: {', '.join(analysts)}")
        content_parts.append(f"**ç ”ç©¶æ·±åº¦**: {research_depth}")
        content_parts.append("")
        content_parts.append("---")
        content_parts.append("")
        
        # æ‰§è¡Œæ‘˜è¦
        if summary:
            content_parts.append("## ğŸ“Š æ‰§è¡Œæ‘˜è¦")
            content_parts.append("")
            content_parts.append(summary)
            content_parts.append("")
            content_parts.append("---")
            content_parts.append("")
        
        # å„æ¨¡å—å†…å®¹
        module_order = [
            "company_overview",
            "financial_analysis", 
            "technical_analysis",
            "market_analysis",
            "risk_analysis",
            "valuation_analysis",
            "investment_recommendation"
        ]
        
        module_titles = {
            "company_overview": "ğŸ¢ å…¬å¸æ¦‚å†µ",
            "financial_analysis": "ğŸ’° è´¢åŠ¡åˆ†æ",
            "technical_analysis": "ğŸ“ˆ æŠ€æœ¯åˆ†æ",
            "market_analysis": "ğŸŒ å¸‚åœºåˆ†æ",
            "risk_analysis": "âš ï¸ é£é™©åˆ†æ",
            "valuation_analysis": "ğŸ’ ä¼°å€¼åˆ†æ",
            "investment_recommendation": "ğŸ¯ æŠ•èµ„å»ºè®®"
        }
        
        # æŒ‰é¡ºåºæ·»åŠ æ¨¡å—
        for module_key in module_order:
            if module_key in reports:
                module_content = reports[module_key]
                if isinstance(module_content, str) and module_content.strip():
                    title = module_titles.get(module_key, module_key)
                    content_parts.append(f"## {title}")
                    content_parts.append("")
                    content_parts.append(module_content)
                    content_parts.append("")
                    content_parts.append("---")
                    content_parts.append("")
        
        # æ·»åŠ å…¶ä»–æœªåˆ—å‡ºçš„æ¨¡å—
        for module_key, module_content in reports.items():
            if module_key not in module_order:
                if isinstance(module_content, str) and module_content.strip():
                    content_parts.append(f"## {module_key}")
                    content_parts.append("")
                    content_parts.append(module_content)
                    content_parts.append("")
                    content_parts.append("---")
                    content_parts.append("")
        
        # é¡µè„š
        content_parts.append("")
        content_parts.append("---")
        content_parts.append("")
        content_parts.append("*æœ¬æŠ¥å‘Šç”± TradingAgents-CN è‡ªåŠ¨ç”Ÿæˆ*")
        content_parts.append("")
        
        markdown_content = "\n".join(content_parts)
        logger.info(f"âœ… Markdown æŠ¥å‘Šç”Ÿæˆå®Œæˆï¼Œé•¿åº¦: {len(markdown_content)} å­—ç¬¦")
        
        return markdown_content
    
    def _clean_markdown_for_pandoc(self, md_content: str) -> str:
        """æ¸…ç† Markdown å†…å®¹ï¼Œé¿å… pandoc è§£æé—®é¢˜"""
        import re

        # ç§»é™¤å¯èƒ½å¯¼è‡´ YAML è§£æé—®é¢˜çš„å†…å®¹
        # å¦‚æœå¼€å¤´æœ‰ "---"ï¼Œåœ¨å‰é¢æ·»åŠ ç©ºè¡Œ
        if md_content.strip().startswith("---"):
            md_content = "\n" + md_content

        # ğŸ”¥ ç§»é™¤å¯èƒ½å¯¼è‡´ç«–æ’çš„ HTML æ ‡ç­¾å’Œæ ·å¼
        # ç§»é™¤ writing-mode ç›¸å…³çš„æ ·å¼
        md_content = re.sub(r'<[^>]*writing-mode[^>]*>', '', md_content, flags=re.IGNORECASE)
        md_content = re.sub(r'<[^>]*text-orientation[^>]*>', '', md_content, flags=re.IGNORECASE)

        # ç§»é™¤ <div> æ ‡ç­¾ä¸­çš„ style å±æ€§ï¼ˆå¯èƒ½åŒ…å«ç«–æ’æ ·å¼ï¼‰
        md_content = re.sub(r'<div\s+style="[^"]*">', '<div>', md_content, flags=re.IGNORECASE)
        md_content = re.sub(r'<span\s+style="[^"]*">', '<span>', md_content, flags=re.IGNORECASE)

        # ğŸ”¥ ç§»é™¤å¯èƒ½å¯¼è‡´é—®é¢˜çš„ HTML æ ‡ç­¾
        # ä¿ç•™åŸºæœ¬çš„ Markdown æ ¼å¼ï¼Œç§»é™¤å¤æ‚çš„ HTML
        md_content = re.sub(r'<style[^>]*>.*?</style>', '', md_content, flags=re.DOTALL | re.IGNORECASE)

        # ğŸ”¥ ç¡®ä¿æ‰€æœ‰æ®µè½éƒ½æ˜¯æ­£å¸¸çš„æ¨ªæ’æ–‡æœ¬
        # åœ¨æ¯ä¸ªæ®µè½å‰åæ·»åŠ æ˜ç¡®çš„æ¢è¡Œï¼Œé¿å… Pandoc è¯¯åˆ¤
        lines = md_content.split('\n')
        cleaned_lines = []
        for line in lines:
            # è·³è¿‡ç©ºè¡Œ
            if not line.strip():
                cleaned_lines.append(line)
                continue

            # å¦‚æœæ˜¯æ ‡é¢˜ã€åˆ—è¡¨ã€è¡¨æ ¼ç­‰ Markdown è¯­æ³•ï¼Œä¿æŒåŸæ ·
            if line.strip().startswith(('#', '-', '*', '|', '>', '```', '1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')):
                cleaned_lines.append(line)
            else:
                # æ™®é€šæ®µè½ï¼šç¡®ä¿æ²¡æœ‰ç‰¹æ®Šå­—ç¬¦å¯¼è‡´ç«–æ’
                cleaned_lines.append(line)

        md_content = '\n'.join(cleaned_lines)

        return md_content

    def _create_pdf_css(self) -> str:
        """åˆ›å»º PDF æ ·å¼è¡¨ï¼Œæ§åˆ¶è¡¨æ ¼åˆ†é¡µå’Œæ–‡æœ¬æ–¹å‘"""
        return """
<style>
/* ğŸ”¥ å¼ºåˆ¶æ‰€æœ‰æ–‡æœ¬æ¨ªæ’æ˜¾ç¤ºï¼ˆä¿®å¤ä¸­æ–‡ç«–æ’é—®é¢˜ï¼‰ */
* {
    writing-mode: horizontal-tb !important;
    text-orientation: mixed !important;
}

body {
    writing-mode: horizontal-tb !important;
    direction: ltr !important;
}

/* æ®µè½å’Œæ–‡æœ¬ */
p, div, span, td, th, li {
    writing-mode: horizontal-tb !important;
    text-orientation: mixed !important;
}

/* è¡¨æ ¼æ ·å¼ - å…è®¸è·¨é¡µ */
table {
    width: 100%;
    border-collapse: collapse;
    page-break-inside: auto;
    writing-mode: horizontal-tb !important;
}

/* è¡¨æ ¼è¡Œ - é¿å…åœ¨è¡Œä¸­é—´åˆ†é¡µ */
tr {
    page-break-inside: avoid;
    page-break-after: auto;
}

/* è¡¨å¤´ - åœ¨æ¯é¡µé‡å¤æ˜¾ç¤º */
thead {
    display: table-header-group;
}

/* è¡¨æ ¼å•å…ƒæ ¼ */
td, th {
    padding: 8px;
    border: 1px solid #ddd;
    writing-mode: horizontal-tb !important;
    text-orientation: mixed !important;
}

/* è¡¨å¤´æ ·å¼ */
th {
    background-color: #f2f2f2;
    font-weight: bold;
}

/* é¿å…æ ‡é¢˜åç«‹å³åˆ†é¡µ */
h1, h2, h3, h4, h5, h6 {
    page-break-after: avoid;
    writing-mode: horizontal-tb !important;
}

/* é¿å…åœ¨åˆ—è¡¨é¡¹ä¸­é—´åˆ†é¡µ */
li {
    page-break-inside: avoid;
}

/* ä»£ç å— */
pre, code {
    writing-mode: horizontal-tb !important;
    white-space: pre-wrap;
    word-wrap: break-word;
}
</style>
"""
    
    def generate_docx_report(self, report_doc: Dict[str, Any]) -> bytes:
        """ç”Ÿæˆ Word æ–‡æ¡£æ ¼å¼æŠ¥å‘Š"""
        logger.info("ğŸ“„ å¼€å§‹ç”Ÿæˆ Word æ–‡æ¡£...")

        if not self.pandoc_available:
            raise Exception("Pandoc ä¸å¯ç”¨ï¼Œæ— æ³•ç”Ÿæˆ Word æ–‡æ¡£ã€‚è¯·å®‰è£… pandoc æˆ–ä½¿ç”¨ Markdown æ ¼å¼å¯¼å‡ºã€‚")

        # ç”Ÿæˆ Markdown å†…å®¹
        md_content = self.generate_markdown_report(report_doc)

        try:
            # åˆ›å»ºä¸´æ—¶æ–‡ä»¶
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
                output_file = tmp_file.name

            logger.info(f"ğŸ“ ä¸´æ—¶æ–‡ä»¶è·¯å¾„: {output_file}")

            # Pandoc å‚æ•°
            extra_args = [
                '--from=markdown-yaml_metadata_block',  # ç¦ç”¨ YAML å…ƒæ•°æ®å—è§£æ
                '--standalone',  # ç”Ÿæˆç‹¬ç«‹æ–‡æ¡£
                '--wrap=preserve',  # ä¿ç•™æ¢è¡Œ
                '--columns=120',  # è®¾ç½®åˆ—å®½
                '-M', 'lang=zh-CN',  # ğŸ”¥ æ˜ç¡®æŒ‡å®šè¯­è¨€ä¸ºç®€ä½“ä¸­æ–‡
                '-M', 'dir=ltr',  # ğŸ”¥ æ˜ç¡®æŒ‡å®šæ–‡æœ¬æ–¹å‘ä¸ºä»å·¦åˆ°å³
            ]

            # æ¸…ç†å†…å®¹
            cleaned_content = self._clean_markdown_for_pandoc(md_content)

            # è½¬æ¢ä¸º Word
            pypandoc.convert_text(
                cleaned_content,
                'docx',
                format='markdown',
                outputfile=output_file,
                extra_args=extra_args
            )

            logger.info("âœ… pypandoc è½¬æ¢å®Œæˆ")

            # ğŸ”¥ åå¤„ç†ï¼šä¿®å¤ Word æ–‡æ¡£ä¸­çš„æ–‡æœ¬æ–¹å‘
            try:
                from docx import Document
                doc = Document(output_file)

                # ä¿®å¤æ‰€æœ‰æ®µè½çš„æ–‡æœ¬æ–¹å‘
                for paragraph in doc.paragraphs:
                    # è®¾ç½®æ®µè½ä¸ºä»å·¦åˆ°å³
                    if paragraph._element.pPr is not None:
                        # ç§»é™¤å¯èƒ½çš„ç«–æ’è®¾ç½®
                        for child in list(paragraph._element.pPr):
                            if 'textDirection' in child.tag or 'bidi' in child.tag:
                                paragraph._element.pPr.remove(child)

                # ä¿®å¤è¡¨æ ¼ä¸­çš„æ–‡æœ¬æ–¹å‘
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                if paragraph._element.pPr is not None:
                                    for child in list(paragraph._element.pPr):
                                        if 'textDirection' in child.tag or 'bidi' in child.tag:
                                            paragraph._element.pPr.remove(child)

                # ä¿å­˜ä¿®å¤åçš„æ–‡æ¡£
                doc.save(output_file)
                logger.info("âœ… Word æ–‡æ¡£æ–‡æœ¬æ–¹å‘ä¿®å¤å®Œæˆ")
            except ImportError:
                logger.warning("âš ï¸ python-docx æœªå®‰è£…ï¼Œè·³è¿‡æ–‡æœ¬æ–¹å‘ä¿®å¤")
            except Exception as e:
                logger.warning(f"âš ï¸ Word æ–‡æ¡£æ–‡æœ¬æ–¹å‘ä¿®å¤å¤±è´¥: {e}")

            # è¯»å–ç”Ÿæˆçš„æ–‡ä»¶
            with open(output_file, 'rb') as f:
                docx_content = f.read()

            logger.info(f"âœ… Word æ–‡æ¡£ç”ŸæˆæˆåŠŸï¼Œå¤§å°: {len(docx_content)} å­—èŠ‚")

            # æ¸…ç†ä¸´æ—¶æ–‡ä»¶
            os.unlink(output_file)

            return docx_content
            
        except Exception as e:
            logger.error(f"âŒ Word æ–‡æ¡£ç”Ÿæˆå¤±è´¥: {e}", exc_info=True)
            # æ¸…ç†ä¸´æ—¶æ–‡ä»¶
            try:
                if 'output_file' in locals() and os.path.exists(output_file):
                    os.unlink(output_file)
            except:
                pass
            raise Exception(f"ç”Ÿæˆ Word æ–‡æ¡£å¤±è´¥: {e}")
    
    def _markdown_to_html(self, md_content: str) -> str:
        """å°† Markdown è½¬æ¢ä¸º HTML"""
        import markdown

        # é…ç½® Markdown æ‰©å±•
        extensions = [
            'markdown.extensions.tables',  # è¡¨æ ¼æ”¯æŒ
            'markdown.extensions.fenced_code',  # ä»£ç å—æ”¯æŒ
            'markdown.extensions.nl2br',  # æ¢è¡Œæ”¯æŒ
        ]

        # è½¬æ¢ä¸º HTML
        html_content = markdown.markdown(md_content, extensions=extensions)

        # æ·»åŠ  HTML æ¨¡æ¿å’Œæ ·å¼
        # WeasyPrint ä¼˜åŒ–çš„ CSSï¼ˆç§»é™¤ä¸æ”¯æŒçš„å±æ€§ï¼‰
        html_template = f"""
<!DOCTYPE html>
<html lang="zh-CN" dir="ltr">
<head>
    <meta charset="UTF-8">
    <title>åˆ†ææŠ¥å‘Š</title>
    <style>
        /* åŸºç¡€æ ·å¼ - ç¡®ä¿æ–‡æœ¬æ–¹å‘æ­£ç¡® */
        html {{
            direction: ltr;
        }}

        body {{
            font-family: "Noto Sans CJK SC", "Microsoft YaHei", "SimHei", "Arial", sans-serif;
            line-height: 1.8;
            color: #333;
            margin: 20mm;
            padding: 0;
            background: white;
            direction: ltr;
        }}

        /* æ ‡é¢˜æ ·å¼ */
        h1, h2, h3, h4, h5, h6 {{
            color: #2c3e50;
            margin-top: 1.5em;
            margin-bottom: 0.8em;
            font-weight: 600;
            page-break-after: avoid;
            direction: ltr;
        }}

        h1 {{
            font-size: 2em;
            border-bottom: 3px solid #3498db;
            padding-bottom: 0.3em;
            page-break-before: always;
        }}

        h1:first-child {{
            page-break-before: avoid;
        }}

        h2 {{
            font-size: 1.6em;
            border-bottom: 2px solid #bdc3c7;
            padding-bottom: 0.25em;
        }}

        h3 {{
            font-size: 1.3em;
            color: #34495e;
        }}

        /* æ®µè½æ ·å¼ */
        p {{
            margin: 0.8em 0;
            text-align: left;
            direction: ltr;
        }}

        /* è¡¨æ ¼æ ·å¼ - ä¼˜åŒ–åˆ†é¡µ */
        table {{
            width: 100%;
            border-collapse: collapse;
            margin: 1.5em 0;
            font-size: 0.9em;
            direction: ltr;
        }}

        /* è¡¨å¤´åœ¨æ¯é¡µé‡å¤ */
        thead {{
            display: table-header-group;
        }}

        tbody {{
            display: table-row-group;
        }}

        /* è¡¨æ ¼è¡Œé¿å…è·¨é¡µæ–­å¼€ */
        tr {{
            page-break-inside: avoid;
        }}

        th, td {{
            border: 1px solid #ddd;
            padding: 10px 12px;
            text-align: left;
            direction: ltr;
        }}

        th {{
            background-color: #3498db;
            color: white;
            font-weight: bold;
        }}

        tbody tr:nth-child(even) {{
            background-color: #f8f9fa;
        }}

        tbody tr:hover {{
            background-color: #e9ecef;
        }}

        /* ä»£ç å—æ ·å¼ */
        code {{
            background-color: #f4f4f4;
            padding: 2px 6px;
            border-radius: 3px;
            font-family: "Consolas", "Monaco", "Courier New", monospace;
            font-size: 0.9em;
            direction: ltr;
        }}

        pre {{
            background-color: #f4f4f4;
            padding: 15px;
            border-radius: 5px;
            border-left: 4px solid #3498db;
            page-break-inside: avoid;
            direction: ltr;
        }}

        pre code {{
            background-color: transparent;
            padding: 0;
        }}

        /* åˆ—è¡¨æ ·å¼ */
        ul, ol {{
            margin: 0.8em 0;
            padding-left: 2em;
            direction: ltr;
        }}

        li {{
            margin: 0.4em 0;
            direction: ltr;
        }}

        /* å¼ºè°ƒæ–‡æœ¬ */
        strong, b {{
            font-weight: 700;
            color: #2c3e50;
        }}

        em, i {{
            font-style: italic;
            color: #555;
        }}

        /* æ°´å¹³çº¿ */
        hr {{
            border: none;
            border-top: 2px solid #ecf0f1;
            margin: 2em 0;
        }}

        /* é“¾æ¥æ ·å¼ */
        a {{
            color: #3498db;
            text-decoration: none;
        }}

        a:hover {{
            text-decoration: underline;
        }}

        /* åˆ†é¡µæ§åˆ¶ */
        @page {{
            size: A4;
            margin: 20mm;

            @top-center {{
                content: "åˆ†ææŠ¥å‘Š";
                font-size: 10pt;
                color: #999;
            }}

            @bottom-right {{
                content: "ç¬¬ " counter(page) " é¡µ";
                font-size: 10pt;
                color: #999;
            }}
        }}

        /* é¿å…å­¤è¡Œå’Œå¯¡è¡Œ */
        p, li {{
            orphans: 3;
            widows: 3;
        }}

        /* å›¾ç‰‡æ ·å¼ */
        img {{
            max-width: 100%;
            height: auto;
            page-break-inside: avoid;
        }}

        /* å¼•ç”¨å—æ ·å¼ */
        blockquote {{
            margin: 1em 0;
            padding: 0.5em 1em;
            border-left: 4px solid #3498db;
            background-color: #f8f9fa;
            font-style: italic;
            page-break-inside: avoid;
        }}
    </style>
</head>
<body>
{html_content}
</body>
</html>
"""
        return html_template

    def _generate_pdf_with_pdfkit(self, html_content: str) -> bytes:
        """ä½¿ç”¨ pdfkit ç”Ÿæˆ PDF"""
        import pdfkit

        logger.info("ğŸ”§ ä½¿ç”¨ pdfkit + wkhtmltopdf ç”Ÿæˆ PDF...")

        # é…ç½®é€‰é¡¹
        options = {
            'encoding': 'UTF-8',
            'enable-local-file-access': None,
            'page-size': 'A4',
            'margin-top': '20mm',
            'margin-right': '20mm',
            'margin-bottom': '20mm',
            'margin-left': '20mm',
        }

        # ç”Ÿæˆ PDF
        pdf_bytes = pdfkit.from_string(html_content, False, options=options)

        logger.info(f"âœ… pdfkit PDF ç”ŸæˆæˆåŠŸï¼Œå¤§å°: {len(pdf_bytes)} å­—èŠ‚")
        return pdf_bytes

    def _generate_pdf_with_xhtml2pdf(self, md_content: str) -> bytes:
        """ä½¿ç”¨ xhtml2pdf ç”Ÿæˆ PDF (è¾“å…¥ Markdownï¼Œä½¿ç”¨ç»å¯¹è·¯å¾„å­—ä½“æ–‡ä»¶)"""
        from xhtml2pdf import pisa
        from io import BytesIO
        import markdown
        import os

        logger.info("ğŸ”§ ä½¿ç”¨ xhtml2pdf ç”Ÿæˆ PDF (Markdown æº)...")
        
        # 1. å¯»æ‰¾å¯ç”¨ä¸­æ–‡å­—ä½“æ–‡ä»¶ç»å¯¹è·¯å¾„
        font_path = None
        # ä¼˜å…ˆåˆ—è¡¨: å®‹ä½“(æœ€é€šç”¨), é»‘ä½“, å¾®è½¯é›…é»‘
        candidates = ['simsun.ttc', 'simhei.ttf', 'msyh.ttc']
        font_dir = r'C:\Windows\Fonts'
        
        for filename in candidates:
            p = os.path.join(font_dir, filename)
            if os.path.exists(p):
                # CSS url() éœ€è¦æ­£æ–œæ 
                font_path = p.replace('\\', '/') 
                logger.info(f"âœ… æ‰¾åˆ°å¯ç”¨çš„ä¸­æ–‡å­—ä½“æ–‡ä»¶: {font_path}")
                break
        
        if not font_path:
            logger.warning("âš ï¸ æœªæ‰¾åˆ°å¸¸ç”¨ä¸­æ–‡å­—ä½“æ–‡ä»¶ï¼ŒPDF ä¸­æ–‡å¯èƒ½ä¹±ç ")
            font_path = "sans-serif" # Fallback

        # 2. Markdown -> HTML
        extensions = [
             'markdown.extensions.tables',
             'markdown.extensions.fenced_code',
             'markdown.extensions.nl2br',
        ]
        raw_html = markdown.markdown(md_content, extensions=extensions)

        # 3. æ„å»º HTML + CSS
        final_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <style>
                @page {{
                    size: A4;
                    margin: 2cm;
                }}
                
                /* å®šä¹‰è‡ªå®šä¹‰å­—ä½“ */
                @font-face {{
                    font-family: 'ChineseFont';
                    src: url('{font_path}');
                }}
                
                body {{
                    font-family: 'ChineseFont', sans-serif;
                    font-size: 10pt;
                    line-height: 1.5;
                }}
                
                /* å¼ºåˆ¶æ‰€æœ‰å…ƒç´ ç»§æ‰¿æˆ–ä½¿ç”¨è¯¥å­—ä½“ */
                h1, h2, h3, h4, h5, h6, p, div, span, li, a, table, th, td, pre, code, b, strong, i, em {{
                    font-family: 'ChineseFont', sans-serif;
                }}

                h1 {{ font-size: 18pt; color: #2c3e50; border-bottom: 2px solid #eee; padding-bottom: 10px; font-weight: bold; }}
                h2 {{ font-size: 16pt; color: #34495e; margin-top: 20px; font-weight: bold; }}
                h3 {{ font-size: 14pt; color: #7f8c8d; font-weight: bold; }}
                p {{ margin-bottom: 10px; text-align: justify; }}
                
                /* è¡¨æ ¼æ ·å¼ */
                table {{
                    width: 100%;
                    border-collapse: collapse; 
                    margin-bottom: 15px;
                }}
                th {{
                    background-color: #f2f2f2;
                    color: #333;
                    font-weight: bold;
                    border: 1px solid #ccc;
                    padding: 5px;
                }}
                td {{
                    border: 1px solid #ccc;
                    padding: 5px;
                }}
                
                /* ä»£ç å— */
                pre {{
                    background-color: #f8f8f8;
                    border: 1px solid #ddd;
                    padding: 10px;
                    white-space: pre-wrap;
                }}
            </style>
        </head>
        <body>
            {raw_html}
        </body>
        </html>
        """

        result_file = BytesIO()
        
        # 4. è½¬æ¢
        pisa_status = pisa.CreatePDF(
            final_html,
            dest=result_file,
            encoding='utf-8'
        )

        if pisa_status.err:
            raise Exception(f"xhtml2pdf è½¬æ¢å†…éƒ¨é”™è¯¯: {pisa_status.err}")

        pdf_value = result_file.getvalue()
        logger.info(f"âœ… xhtml2pdf PDF ç”ŸæˆæˆåŠŸï¼Œå¤§å°: {len(pdf_value)} å­—èŠ‚")
        return pdf_value

    def generate_pdf_report(self, report_doc: Dict[str, Any]) -> bytes:
        """ç”Ÿæˆ PDF æ ¼å¼æŠ¥å‘Š"""
        logger.info("ğŸ“Š å¼€å§‹ç”Ÿæˆ PDF æ–‡æ¡£...")

        # æ£€æŸ¥ pdfkit æ˜¯å¦å¯ç”¨
        if not self.pdfkit_available:
            xhtml2pdf_error = None
            if self.xhtml2pdf_available:
                logger.info("âš ï¸ pdfkit ä¸å¯ç”¨ï¼Œå°è¯•åˆ‡æ¢ä½¿ç”¨ xhtml2pdf...")
                try:
                    md_content = self.generate_markdown_report(report_doc)
                    return self._generate_pdf_with_xhtml2pdf(md_content)
                except Exception as e:
                     xhtml2pdf_error = str(e)
                     logger.error(f"âŒ xhtml2pdf ç”Ÿæˆå¤±è´¥: {e}")
            
            error_msg = (
                "PDF ç”Ÿæˆå·¥å…·ä¸å¯ç”¨ã€‚\n"
                "è¯·å®‰è£… wkhtmltopdf (æ¨è) æˆ–æ£€æŸ¥åå°æ—¥å¿—ã€‚\n"
                f"pdfkit çŠ¶æ€: {PDFKIT_ERROR or 'æœªå®‰è£…/æœªæ‰¾åˆ° executable'}\n"
                f"xhtml2pdf çŠ¶æ€: {'å¯ç”¨' if self.xhtml2pdf_available else 'æœªå®‰è£…'} \n"
                f"xhtml2pdf é”™è¯¯: {xhtml2pdf_error or 'æœªå°è¯• (ä¾èµ–ç¼ºå¤±)'}"
            )
            logger.error(f"âŒ {error_msg}")
            raise Exception(error_msg)

        # ç”Ÿæˆ Markdown å†…å®¹
        md_content = self.generate_markdown_report(report_doc)

        # ä½¿ç”¨ pdfkit ç”Ÿæˆ PDF
        try:
            html_content = self._markdown_to_html(md_content)
            return self._generate_pdf_with_pdfkit(html_content)
        except Exception as e:
            logger.warning(f"âš ï¸ pdfkit ç”Ÿæˆå¤±è´¥: {e}ï¼Œå°è¯•ä½¿ç”¨ xhtml2pdf é™çº§å¤„ç†...")
            xhtml2pdf_error_fallback = None
            if self.xhtml2pdf_available:
                try:
                    return self._generate_pdf_with_xhtml2pdf(md_content)
                except Exception as inner_e:
                    xhtml2pdf_error_fallback = str(inner_e)
                    logger.error(f"âŒ xhtml2pdf é™çº§å°è¯•ä¹Ÿå¤±è´¥: {inner_e}")
            
            error_msg = f"PDF ç”Ÿæˆå¤±è´¥: {e}"
            if xhtml2pdf_error_fallback:
                error_msg += f" (é™çº§é‡è¯•ä¹Ÿå¤±è´¥: {xhtml2pdf_error_fallback})"
            
            logger.error(f"âŒ {error_msg}")
            raise Exception(error_msg)


# åˆ›å»ºå…¨å±€å¯¼å‡ºå™¨å®ä¾‹
report_exporter = ReportExporter()

